
from docx.shared import Inches
from docx import Document
import os
import tempfile


def report_combined_content(dict):
    print(dict)
    #get infor details from dictionary
    name = dict["name"]
    company = dict["company"].capitalize()
    person_debt = dict["person_debt"]
    business_debt= dict["business_debt"]
    person_income= dict["person_income"]
    business_income= dict["business_income"]
    DTO= dict["DTO"]
    DTI= dict["DTI"]
    GDSCR= dict["GDSCR"]
    capacity= dict["capacity"]
    balance= dict["balance"]
    ex= dict["ex"]
    eq = dict["eq"]
    tu = dict["tu"]
    persons_count = len(name)
    loan_limit = dict['loan_monthly_limit']
    requested_funds = dict['requested_funds']
    prime_rate = dict['prime_rate'] 
    additional_rate = dict['additional_rate'] 
    Total_interest = int(prime_rate) + int(additional_rate)
    tenure = dict['tenure'] 
    estimated_monthly_payment = dict['estimated_monthly_payment'] 
    estimated_gdscr = dict['estimated_gdscr']
    is_eligible = dict['is_eligible']
    
    #Calculation
    try:
        total_income = int(person_income) + int(business_income)
    except ValueError:
        total_income = 0
    try:

        total_debt_pmt = int(person_debt) + int(business_debt)
    except ValueError:
        total_debt_pmt = 0

    #Change to string to be able to concate in the documents
    total_income = str(total_income)
    total_debt_pmt = str(total_debt_pmt)
    GDSCR = str(GDSCR)
    DTO = str(DTO)
    DTI = str(DTI)
    balance = str(balance)
    capacity= str(capacity)
    
    #Initiate document
    document = Document()

    #Header
    document.add_picture('static/images/wcf_logo.png', width=Inches(3))
    document.add_heading('Personal Business Financial Analysis ', 0)
    document.add_heading(company, 0)
    
    document.add_paragraph('Prepared by Wallace Capital Funding, LLC ')
    

    #Personal Credit Profile
    document.add_heading('Personal Credit Profile', level=1)
    document.add_paragraph('Credit Scores', style='List Bullet')

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'TU'
    hdr_cells[2].text = 'EQ'
    hdr_cells[3].text = 'EX'

    for k in range(0, persons_count):
        row_cells = table.add_row().cells
        row_cells[0].text = name[k]
        row_cells[1].text = str(ex[k])
        row_cells[2].text = str(eq[k])
        row_cells[3].text = str(tu[k])

    all_names = ', '.join(name)
    document.add_paragraph('PCW: '+ all_names, style='List Bullet')
    document.add_paragraph('   Debt-to-Owing Ratio (DTO): '+DTO)
    document.add_paragraph('   Debt-to-Income Ratio (DTI): '+DTI)
    
    #Cashflow analysis
    document.add_heading('Cash Flow Analysis ', level=1)
    document.add_paragraph('GDSCR', style='List Bullet')
    document.add_paragraph("   The Global Cash Flow Sheet allows us to determine the total operating income and total debt. We compute the operating income using information from the company's balance sheet, profit and loss statement, and tax returns. The total debt is calculated based on recent credit report statements. The average net operating income is $" + total_income + " per month. The total monthly debt payments is $"+ total_debt_pmt +". This results in a Gross Debt Service Coverage Ratio (GDSCR) of " + GDSCR + ".")
    
    #Bank Statement Analsyis
    document.add_heading('Bank Statement Analysis', level=1)
    document.add_paragraph(" We conducted a bank statement analysis and the analysis showed that the average monthly balance is "+ balance +". While based on your average deposit and withdrawal, you have a capacity of "+ capacity +"." , style='List Bullet')
    loan_monthly_limit =  loan_limit   #calculate_loan_limit(float(total_income), float(total_debt_pmt))
    
    # document.add_paragraph(, style='List Bullet')
    repayment_months = int(tenure) * 12
    is_eligible = 1
    if is_eligible == 1:
        analysis_str = "You are eligible for this LOAN. "
    else:
        analysis_str = "You are not eligible for this LOAN. "

    req_string = "{} requested funds ${} with total interest rate {}% and repayment tenure is {} years.({} months)".format(company, str(requested_funds), str(Total_interest), tenure, repayment_months)
    analysis_str = analysis_str + "Your monthly loan limit which can be handled is ${}. Your estimated monthly payment for the requested amount ${} is ${} for {} months and estimated GDSCR is {}.".format( str(loan_monthly_limit),  str(requested_funds), str(estimated_monthly_payment), repayment_months , estimated_gdscr)

    document.add_heading('Requested Loan Information', level=1)
    document.add_paragraph(req_string, style='List Bullet')
    
    document.add_heading('Loan Analysis Information', level=1)
    document.add_paragraph(analysis_str, style='List Bullet')
    
    # document.add_paragraph("Average Monthly Income :" + str(total_income) , style='List Bullet')
    # document.add_paragraph("Average Monthly Debt : " + str(total_debt_pmt) , style='List Bullet')
    # document.add_paragraph("Estimated Monthly Payment :" + str(estimated_monthly_payment) , style='List Bullet')
    # document.add_paragraph("Estimated Gdscr :" + str(estimated_gdscr) , style='List Bullet')
    # document.add_paragraph("Loan Monthly limit which can be handled : " + str(loan_monthly_limit) , style='List Bullet')
    # document.add_paragraph("Total Requested Loan amount : " + str(requested_funds), style='List Bullet')
    # document.add_paragraph("Interest Rate: " + str(Total_interest), style='List Bullet')
    # document.add_paragraph("Repayment Tenure: " + str(tenure), style='List Bullet')
    
    #Recommendations
    document.add_heading('Recommendation', level=1)
    document.add_paragraph("   Write recomendation here", style='List Bullet')
    
    return document

def tmp_dir():
    with tempfile.TemporaryDirectory() as tmpdirname:
        return(tempfile.gettempdir())

def generate_combined_report(dict):

    #get name
    name = dict["company"]

    #Created temporary directory
    directory = tmp_dir()

    filename = "Combined BFA for " + name + ".docx"
    file_path = os.path.join(directory, filename)

    #Generate Report
    report = report_combined_content(dict)

    #Save the report in temporary directory
    report.save(file_path)
    
    return file_path